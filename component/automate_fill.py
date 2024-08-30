# import json
# from docx import Document
# from docx2pdf import convert

# def fill(template_path, output_path, data):
#     # Convert "Ngày sinh" into day, month, year if present in data
#     if "Ngày sinh" in data:
#         ngay_sinh = data["Ngày sinh"]
#         ngay, thang, nam = ngay_sinh.split('/')
#         data["Ngày"] = ngay
#         data["Tháng"] = thang
#         data["Năm"] = nam
    
#     # Convert "Họ và tên" to uppercase if present in data
#     if "Họ và tên" in data:
#         data["Họ và tên"] = data["Họ và tên"].upper()
    
#     doc = Document(template_path)

#     # Iterate over each paragraph in the document
#     for paragraph in doc.paragraphs:
#         for key, value in data.items():
#             placeholder = f"[{key}]"
#             if placeholder in paragraph.text:
#                 # Iterate over each run in the paragraph to replace placeholders
#                 for run in paragraph.runs:
#                     if placeholder in run.text:
#                         run.text = run.text.replace(placeholder, value)
    
#     # Save the filled document
#     doc.save(output_path)

#     # Convert the Word document to PDF
#     convert(output_path)
import json
from docx import Document
from docx2pdf import convert
import os

def fill(template_path, output_dir, data):
    # Convert "Ngày sinh" into day, month, year if present in data
    if "Ngày sinh" in data:
        ngay_sinh = data["Ngày sinh"]
        ngay, thang, nam = ngay_sinh.split('/')
        data["Ngày"] = ngay
        data["Tháng"] = thang
        data["Năm"] = nam
    
    # Convert "Họ và tên" to uppercase if present in data
    if "Họ và tên" in data:
        data["Họ và tên"] = data["Họ và tên"].upper()

    # Generate the output file name based on "Họ và tên" or other data fields
    file_name = data.get("Họ và tên", "output").replace(" ", "_") + ".docx"
    output_path = os.path.join(output_dir, file_name)
    
    doc = Document(template_path)

    # Iterate over each paragraph in the document
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            placeholder = f"[{key}]"
            if placeholder in paragraph.text:
                # Iterate over each run in the paragraph to replace placeholders
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)
    
    # Save the filled document
    doc.save(output_path)

    # Convert the Word document to PDF
    pdf_path = output_path.replace(".docx", ".pdf")
    convert(output_path, pdf_path)

    # Xóa file .docx sau khi đã chuyển đổi sang PDF
    if os.path.exists(output_path):
        os.remove(output_path)
    
    return pdf_path
