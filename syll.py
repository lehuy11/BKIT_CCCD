from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def export_to_docx(data, file_name='biểu_mẫu_thông_tin.docx'):
    """
    Xuất dữ liệu ra tệp DOCX.
    :param data: Dữ liệu cần xuất (dạng dictionary hoặc object)
    :param file_name: Tên tệp đầu ra (mặc định: biểu_mẫu_thông_tin.docx)
    """
    # Tạo một tài liệu Word mới
    doc = Document()

    # Thêm tiêu đề chính
    title = doc.add_heading('Biểu Mẫu Thông Tin', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Thêm đoạn giới thiệu
    doc.add_paragraph('Thông tin cá nhân được cung cấp như sau:')

    # Thêm dữ liệu vào tài liệu
    doc.add_paragraph(f"Họ và tên: {data.get('name', 'Không có dữ liệu')}")
    doc.add_paragraph(f"Ngày sinh: {data.get('dob', 'Không có dữ liệu')}")
    doc.add_paragraph(f"CMND/CCCD: {data.get('id', 'Không có dữ liệu')}")
    doc.add_paragraph(f"Địa chỉ: {data.get('address', 'Không có dữ liệu')}")

    # Thêm một khoảng trắng
    doc.add_paragraph()

    # Thêm bảng (nếu cần)
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Thông tin'
    hdr_cells[1].text = 'Chi tiết'

    # Định dạng tiêu đề bảng
    hdr_cells[0].paragraphs[0].runs[0].bold = True
    hdr_cells[1].paragraphs[0].runs[0].bold = True

    for key, value in data.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(value)

    # Thêm đoạn kết thúc
    doc.add_paragraph('\nCảm ơn bạn đã cung cấp thông tin.')

    # Lưu tài liệu
    doc.save(file_name)
    print(f"Biểu mẫu đã được lưu thành công tại: {file_name}")

# Ví dụ sử dụng hàm xuất biểu mẫu
data = {
    'name': 'Nguyễn Văn A',
    'dob': '01/01/1990',
    'id': '0123456789',
    'address': '123 Đường ABC, Quận 1, TP.HCM',
    'phone': '0987654321',
    'email': 'nguyenvana@example.com'
}

# Gọi hàm để xuất biểu mẫu ra file DOCX
export_to_docx(data)
